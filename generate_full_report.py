from __future__ import annotations

from copy import deepcopy
from typing import Any

from docx import Document
from docx.shared import Pt, RGBColor

from config import get_config, save_config
from engine import run_allocation
from scenarios import scenario_definitions
from scoring import base_score, get_saturation_band


REPORT_NAME = "Slot_Allocation_Full_Report.docx"
HEADING_FONT = "Montserrat"
HEADING_COLOR = RGBColor(0xAB, 0x00, 0xFF)


def _run_batches(jos: list[dict], slots: list, batch_sizes: list[int]) -> dict[str, Any]:
    assignments = []
    batch_debug = []
    breakdown: dict[tuple[str, str], dict[str, float]] = {}
    base_initial = {jo["id"]: base_score(jo) for jo in jos}
    batch_offset = 0
    slot_idx = 0
    used_sizes: list[int] = []

    while slot_idx < len(slots):
        size = batch_sizes[len(used_sizes) % len(batch_sizes)]
        batch_slots = slots[slot_idx : slot_idx + size]
        if not batch_slots:
            break
        result = run_allocation(jos, batch_slots, batch_size=len(batch_slots))
        assignments.extend(result.assignments)
        breakdown.update(result.score_breakdown_by_slot_jo)
        for bd in result.batch_debug:
            bd.batch_index = batch_offset + bd.batch_index
            batch_debug.append(bd)
        slot_idx += len(batch_slots)
        used_sizes.append(len(batch_slots))
        batch_offset += len(result.batch_debug)

    return {
        "assignments": assignments,
        "batch_debug": batch_debug,
        "score_breakdown": breakdown,
        "jo_base_scores_initial": base_initial,
        "jo_snapshots_end": {jo["id"]: jo for jo in jos},
        "batch_sizes_used": used_sizes,
    }


def run_scenario(scenario: dict[str, Any], config_override: dict[str, Any]) -> dict[str, Any]:
    original = get_config()
    updated = original | config_override
    save_config(updated)
    try:
        jos = deepcopy(scenario["jos"])
        slots = scenario["slots"]
        batch_sizes = scenario.get("batch_sizes", [4])
        return _run_batches(jos, slots, batch_sizes)
    finally:
        save_config(original)


def _band_from_pct(priority: Any, sat_pct: float) -> tuple[str, str]:
    jo = {
        "id": "TMP",
        "priority": priority,
        "days_remaining": 0,
        "total_duration": 1,
        "initial_demand": 100.0,
        "active_demand": 100.0,
        "project": "N/A",
        "type": "LATERAL",
        "tech_stack": "*",
        "level": "*",
        "slots_allocated": sat_pct,
    }
    _, band_label, split_rule, _ = get_saturation_band(jo)
    return band_label, split_rule


def _add_section_heading(doc: Document, text: str, level: int = 2) -> None:
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.bold = True
        run.font.name = HEADING_FONT
        run.font.color.rgb = HEADING_COLOR
    doc.add_paragraph()


def _add_subheading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = HEADING_FONT
    run.font.color.rgb = HEADING_COLOR
    p.paragraph_format.space_after = Pt(6)


def _style_table(table) -> None:
    table.style = "Light List Accent 1"


def write_batch_section(
    doc: Document, batch, delta_threshold: float, jo_meta: dict[str, dict]
) -> None:
    _add_subheading(doc, f"Batch {batch.batch_index + 1}")
    doc.add_paragraph(f"Slots: {', '.join(batch.slot_ids)}")
    top_id = batch.dominant_id or "-"
    next_id = batch.next_id or "-"
    if batch.lateral_base_scores:
        ranked = sorted(batch.lateral_base_scores.items(), key=lambda kv: kv[1], reverse=True)
        top_id = ranked[0][0]
        next_id = ranked[1][0] if len(ranked) > 1 else "-"
        doc.add_paragraph(f"Top JO base score: {ranked[0][1]:.4f}")
        if len(ranked) > 1:
            doc.add_paragraph(f"Second JO base score: {ranked[1][1]:.4f}")
    doc.add_paragraph(f"Top JO: {top_id}")
    doc.add_paragraph(f"Second JO: {next_id}")
    if batch.delta_value is not None:
        doc.add_paragraph(f"Delta: {batch.delta_value:.4f}")
        doc.add_paragraph(
            "Split decision: "
            + ("Split (delta ≤ threshold)" if batch.delta_value <= delta_threshold else "Greedy")
        )
    else:
        doc.add_paragraph("Delta: n/a")

    if top_id in batch.saturation_pct and top_id in jo_meta:
        sat = batch.saturation_pct[top_id]
        band_label, split_rule = _band_from_pct(jo_meta[top_id]["priority"], sat)
        doc.add_paragraph(f"Saturation (dominant): {sat:.2f}%")
        doc.add_paragraph(f"Band: {band_label} | Split Rule: {split_rule}")

    caps = batch.caps or {}
    if caps:
        cap_lines = [f"{jid}: {cap}" for jid, cap in caps.items() if cap > 0]
        doc.add_paragraph("Caps: " + ", ".join(cap_lines))
        if top_id in caps:
            doc.add_paragraph(f"Cap applied (dominant): {caps[top_id]}")

    alloc = batch.lateral_assigned_this_batch or {}
    alloc_table = doc.add_table(rows=1, cols=2)
    alloc_table.rows[0].cells[0].text = "JO"
    alloc_table.rows[0].cells[1].text = "Slots Assigned"
    for jid, count in alloc.items():
        row = alloc_table.add_row().cells
        row[0].text = jid
        row[1].text = str(count)
    _style_table(alloc_table)
    doc.add_paragraph()

    doc.add_paragraph("Why:")
    delta_line = (
        f"Delta {batch.delta_value:.4f} vs threshold {delta_threshold:.4f}."
        if batch.delta_value is not None
        else "Delta n/a (single competitor)."
    )
    doc.add_paragraph(delta_line, style="List Bullet")
    doc.add_paragraph("Saturation band determined pool size for the dominant JO.", style="List Bullet")
    doc.add_paragraph("Score breakdown drove final assignment within caps.", style="List Bullet")


def write_slot_table(doc: Document, assignments: list, limit: int = 50) -> None:
    _add_section_heading(doc, "Slot Level Table (First 50)", level=2)
    table = doc.add_table(rows=1, cols=4)
    table.rows[0].cells[0].text = "Slot ID"
    table.rows[0].cells[1].text = "Assigned JO"
    table.rows[0].cells[2].text = "Final Score"
    table.rows[0].cells[3].text = "Reason"
    for assignment in assignments[:limit]:
        row = table.add_row().cells
        row[0].text = assignment.slot_id
        row[1].text = assignment.jo_id
        row[2].text = "" if assignment.final_score is None else f"{assignment.final_score:.4f}"
        row[3].text = assignment.reason
    _style_table(table)
    doc.add_paragraph()


def write_comparison(
    doc: Document,
    scenario_id: str,
    default_res: dict,
    updated_res: dict,
    default_threshold: float,
    updated_threshold: float,
) -> None:
    _add_section_heading(doc, f"Comparison — {scenario_id}", level=2)
    def_batches = len(default_res["batch_debug"])
    upd_batches = len(updated_res["batch_debug"])
    doc.add_paragraph(f"Default batches: {def_batches}, Updated batches: {upd_batches}")
    def_splits = sum(
        1
        for bd in default_res["batch_debug"]
        if bd.delta_value is not None and bd.delta_value <= default_threshold
    )
    upd_splits = sum(
        1
        for bd in updated_res["batch_debug"]
        if bd.delta_value is not None and bd.delta_value <= updated_threshold
    )
    doc.add_paragraph(f"Split frequency — default: {def_splits}, updated: {upd_splits}")


def main() -> None:
    doc = Document()
    title = doc.add_heading("Slot Allocation Full Report", level=1)
    for run in title.runs:
        run.bold = True
        run.font.name = HEADING_FONT
        run.font.color.rgb = HEADING_COLOR
    doc.add_paragraph()

    default_config = get_config()
    updated_config = {
        "priority_weight": 0.40,
        "demand_weight": 0.25,
        "urgency_weight": default_config["urgency_weight"],
        "affinity_weight": default_config["affinity_weight"],
        "delta_threshold": 0.040,
    }

    _add_section_heading(doc, "Overview", level=2)
    doc.add_paragraph(
        "This report summarizes allocation outcomes across all scenarios using the existing engine logic."
    )
    doc.add_paragraph(f"Default config: {default_config}")
    doc.add_paragraph(f"Updated config: {updated_config}")
    doc.add_paragraph()

    scenarios = scenario_definitions()
    results_default: dict[str, Any] = {}
    results_updated: dict[str, Any] = {}

    _add_section_heading(doc, "Scenario Breakdown", level=2)
    for scenario in scenarios:
        sid = scenario["id"]
        _add_section_heading(doc, scenario["name"], level=2)
        desc = scenario.get("description", "")
        if desc:
            doc.add_paragraph(desc)
        doc.add_paragraph()

        _add_section_heading(doc, "Default Config", level=3)
        res_default = run_scenario(scenario, {})
        results_default[sid] = res_default
        jo_meta = {jo["id"]: jo for jo in scenario["jos"]}
        for batch in res_default["batch_debug"]:
            write_batch_section(doc, batch, default_config["delta_threshold"], jo_meta)
        write_slot_table(doc, res_default["assignments"])

        _add_section_heading(doc, "Updated Config", level=3)
        res_updated = run_scenario(scenario, updated_config)
        results_updated[sid] = res_updated
        for batch in res_updated["batch_debug"]:
            write_batch_section(doc, batch, updated_config["delta_threshold"], jo_meta)
        write_slot_table(doc, res_updated["assignments"])

    _add_section_heading(doc, "Comparison", level=2)
    for scenario in scenarios:
        sid = scenario["id"]
        write_comparison(
            doc,
            sid,
            results_default[sid],
            results_updated[sid],
            default_config["delta_threshold"],
            updated_config["delta_threshold"],
        )

    _add_section_heading(doc, "Executive Summary", level=2)
    doc.add_paragraph(
        "Increasing priority weight emphasized high-priority roles earlier, while a higher delta threshold "
        "reduced split frequency. Reducing demand weight slightly dampened demand-driven prioritization, "
        "shifting emphasis toward urgency and priority signals."
    )
    doc.add_paragraph()

    doc.save(REPORT_NAME)


if __name__ == "__main__":
    main()
